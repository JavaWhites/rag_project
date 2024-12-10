import io
import os
import re
import subprocess
import docx
from PIL import Image
from docx import Document


def get_outline_level(input_xml):
    """
    从xml字段中提取大纲等级数字。

    参数:
        input_xml (str): XML字符串。

    返回:
        int: 大纲等级数字。
    """
    start_index = input_xml.find("<w:outlineLvl")
    end_index = input_xml.find(">", start_index)
    number = input_xml[start_index: end_index + 1]
    return int(re.search(r"\d+", number).group())


def is_title(paragraph):
    """
    判断段落是否设置了大纲等级。

    参数:
        paragraph (docx.text.paragraph.Paragraph): 段落对象。

    返回:
        int or None: 大纲等级或None。
    """
    if paragraph.text.strip() == "":
        return None

    paragraph_xml = paragraph._p.xml
    if "<w:outlineLvl" in paragraph_xml:
        return get_outline_level(paragraph_xml)

    target_style = paragraph.style
    while target_style is not None:
        if "<w:outlineLvl" in target_style.element.xml:
            return get_outline_level(target_style.element.xml)
        target_style = target_style.base_style

    return None


def convert_to_json_structure(document):
    """
    将文档结构转换为JSON格式。

    参数:
        document (docx.Document): 文档对象。

    返回:
        list: JSON格式的文档结构。
    """
    result = []
    level_0_index = 0
    level_1_index = 0
    level_2_index = 0
    level_3_index = 0
    level_4_index = 0

    current_level_0 = None
    current_level_1 = None
    current_level_2 = None
    current_level_3 = None

    for paragraph in document.paragraphs:
        num = is_title(paragraph)
        if num == 0:
            current_level_0 = {"key": f"{level_0_index}", "label": paragraph.text}
            result.append(current_level_0)
            level_0_index += 1
            level_1_index = 0
        elif num == 1 and current_level_0 is not None:
            current_level_1 = {
                "key": f'{current_level_0["key"]}-{level_1_index}',
                "label": paragraph.text,
            }
            if "children" not in current_level_0:
                current_level_0["children"] = []
            current_level_0["children"].append(current_level_1)
            level_1_index += 1
            level_2_index = 0
        elif num == 2 and current_level_1 is not None:
            current_level_2 = {
                "key": f'{current_level_1["key"]}-{level_2_index}',
                "label": paragraph.text,
            }
            if "children" not in current_level_1:
                current_level_1["children"] = []
            current_level_1["children"].append(current_level_2)
            level_2_index += 1
            level_3_index = 0
        elif num == 3 and current_level_2 is not None:
            current_level_3 = {
                "key": f'{current_level_2["key"]}-{level_3_index}',
                "label": paragraph.text,
            }
            if "children" not in current_level_2:
                current_level_2["children"] = []
            current_level_2["children"].append(current_level_3)
            level_3_index += 1
            level_4_index = 0
        elif num == 4 and current_level_3 is not None:
            current_level_4 = {
                "key": f'{current_level_3["key"]}-{level_4_index}',
                "label": paragraph.text,
            }
            if "children" not in current_level_3:
                current_level_3["children"] = []
            current_level_3["children"].append(current_level_4)
            level_4_index += 1
    return result


# def convert_to_json_structure(document):
#     """
#     将文档结构转换为JSON格式。

#     参数:
#         document (docx.Document): 文档对象。

#     返回:
#         list: JSON格式的文档结构。
#     """
#     result = []
#     level_0_index = 0
#     level_1_index = 0
#     level_2_index = 0


#     current_level_0 = None
#     current_level_1 = None

#     for paragraph in document.paragraphs:
#         num = is_title(paragraph)
#         if num == 0:
#             current_level_0 = {"key": f"{level_0_index}", "label": paragraph.text}
#             result.append(current_level_0)
#             level_0_index += 1
#             level_1_index = 0
#         elif num == 1 and current_level_0 is not None:
#             current_level_1 = {
#                 "key": f'{current_level_0["key"]}-{level_1_index}',
#                 "label": paragraph.text,
#             }
#             if "children" not in current_level_0:
#                 current_level_0["children"] = []
#             current_level_0["children"].append(current_level_1)
#             level_1_index += 1
#             level_2_index = 0
#         elif num == 2 and current_level_1 is not None:
#             current_level_2 = {
#                 "key": f'{current_level_1["key"]}-{level_2_index}',
#                 "label": paragraph.text,
#             }
#             if "children" not in current_level_1:
#                 current_level_1["children"] = []
#             current_level_1["children"].append(current_level_2)
#             level_2_index += 1

#     return result


def convert_emf_to_image(emf_blob):
    """
    将 EMF 二进制数据转换为 PNG 格式。

    参数:
        emf_blob (bytes): EMF 文件的二进制数据。

    返回:
        bytes: 转换后的 PNG 图像的二进制数据。
    """
    # 将 EMF 数据写入临时文件
    emf_filename = "temp.emf"
    png_filename = "temp.png"

    with open(emf_filename, "wb") as f:
        f.write(emf_blob)

    try:
        # 使用 Inkscape 将 EMF 文件转换为 PNG
        subprocess.run(
            ["inkscape", emf_filename, "--export-type=png", "-o", png_filename],
            check=True,
        )

        # 读取 PNG 文件
        with open(png_filename, "rb") as f:
            png_data = f.read()

    finally:
        # 清理临时文件
        if os.path.exists(emf_filename):
            os.remove(emf_filename)
        if os.path.exists(png_filename):
            os.remove(png_filename)

    return png_data


def get_picture(document, element):
    """
    从元素中提取图片，包括普通图片和嵌入对象中的图片。

    参数:
        document (docx.Document): 文档对象。
        element (xml.etree.ElementTree.Element): XML元素。

    返回:
        list: 图片对象列表。
    """
    images = []

    # 定义命名空间
    namespaces = {
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
        "v": "urn:schemas-microsoft-com:vml",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    }

    # 查找普通图片
    img = element.findall(".//pic:pic", namespaces)
    for i in img:
        embed = i.find(".//a:blip", namespaces).get(
            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
        )
        related_part = document.part.related_parts[embed]
        images.append(related_part.blob)

    # 查找嵌入对象中的图片
    embedded_imgs = element.findall(".//v:imagedata", namespaces)
    for img in embedded_imgs:
        embed = img.get(
            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"
        )
        if embed:
            related_part = document.part.related_parts[embed]
            images.append(convert_emf_to_image(related_part.blob))

    return images if images else None


'''
 if element.tag.endswith("p"):
            paragraph = docx.text.paragraph.Paragraph(element, document)
            title_level = is_title(paragraph)

            # 检查当前段落是否是目标章节标题
            if paragraph.text.strip() == chapter_title:
                if chapter_title_level is None:
                    chapter_title_level = title_level
                in_target_chapter = True
                new_doc.add_heading(paragraph.text, level=title_level)
                continue  # 继续检查下一个段落

            # 如果处于目标章节中，检查标题级别并处理内容
            if in_target_chapter:
                if title_level is not None and title_level <= chapter_title_level:
                    # 遇到同级或更高的标题，结束提取
                    break
                else:
                    images = get_picture(document, element)
                    if images:
                        for image in images:
                            new_doc.add_picture(io.BytesIO(image))
                    else:
                        if title_level:
                            new_doc.add_heading(paragraph.text, level=title_level)
                        else:
                            new_doc.add_paragraph(paragraph.text)

'''


def convert_wmf_to_png(wmf_data):
    with open('temp.wmf', 'wb') as temp_wmf:
        temp_wmf.write(wmf_data)

    png_path = 'temp.png'
    # 调用 ImageMagick 将 WMF 转换为 PNG
    subprocess.run(['magick', 'convert', 'temp.wmf', png_path])

    with open(png_path, 'rb') as png_file:
        png_data = png_file.read()

    # 清理临时文件
    subprocess.run(['rm', 'temp.wmf', png_path])
    return png_data


def extract_chapter_to_new_docx(document, chapter_title, output_filename):
    """
    提取指定章节并保存为新的Word文档。

    参数:
        document (docx.Document): 原始文档对象。
        chapter_title (str): 章节标题。
        output_filename (str): 输出文件名。
    """
    new_doc = docx.Document()
    in_target_chapter = False
    chapter_title_level = None

    for element in document.element.body:
        if element.tag.endswith("p"):
            paragraph = docx.text.paragraph.Paragraph(element, document)
            title_level = is_title(paragraph)

            if paragraph.text.strip() == chapter_title:
                chapter_title_level = title_level

            if chapter_title_level is None:
                continue

            if (
                    title_level == chapter_title_level
                    and paragraph.text.strip() == chapter_title
            ):
                in_target_chapter = True
                new_doc.add_heading(paragraph.text, level=title_level)

            elif (
                    title_level is not None
                    and title_level <= chapter_title_level
                    and in_target_chapter
            ):
                break
            elif in_target_chapter:
                images = get_picture(document, element)
                if images:
                    for image in images:
                        image_stream = io.BytesIO(image)
                        try:
                            img = Image.open(image_stream)
                            img_format = img.format

                            # 如果是 WMF 格式，转换为 PNG
                            if img_format == 'WMF':
                                png_data = convert_wmf_to_png(image)
                                new_doc.add_picture(io.BytesIO(png_data))
                            else:
                                img = img.convert("RGB")  # 转换为支持的格式
                                img_stream = io.BytesIO()
                                img.save(img_stream, format='PNG')  # 保存为 PNG 格式
                                img_stream.seek(0)  # 重置流位置
                                new_doc.add_picture(img_stream)  # 添加图片

                        except Exception as e:
                            print(f"Error processing image: {e}")

                        # new_doc.add_picture(io.BytesIO(image))
                else:
                    if title_level:
                        new_doc.add_heading(paragraph.text, level=title_level)
                    else:
                        new_doc.add_paragraph(paragraph.text)

        elif element.tag.endswith("tbl") and in_target_chapter:
            tbl = docx.table.Table(element, document)
            new_tbl = new_doc.add_table(rows=0, cols=len(tbl.columns))
            for row in tbl.rows:
                new_row = new_tbl.add_row().cells
                for idx, cell in enumerate(row.cells):
                    # new_row[idx].text = cell.text
                    paragraphs = cell.paragraphs.copy()
                    for paragraph in paragraphs:
                        images = get_picture(document, paragraph._element)
                        if images:
                            for image in images:
                                new_row[idx].add_paragraph().add_run().add_picture(
                                    io.BytesIO(image)
                                )
                        else:
                            new_row[idx].add_paragraph(paragraph.text)

    new_doc.save(output_filename)


# def read_title(path):
#     document = docx.Document(path)

#     # 转换为JSON结构
#     json_structure = convert_to_json_structure(document)

#     case_names = []
#     for d in json_structure:
#         if d["label"] == "工艺设计详细方案":
#             for dc in d["children"]:
#                 if dc["label"] == "业务需求":
#                     for c in dc["children"]:
#                         case_names.append(c["label"])

#     return case_names
def read_title(path):
    document = docx.Document(path)

    # 转换为JSON结构
    json_structure = convert_to_json_structure(document)

    case_names_one = []
    case_names_two = []
    case_names_three = []
    case_names_four = []
    case_names_five = []
    for d1 in json_structure:
        if d1["label"] != "":
            case_names_one.append(d1['label'])
            if 'children' in d1:
                for d2 in d1['children']:
                    if d2['label'] != "":
                        case_names_two.append(d2['label'])
                        if 'children' in d2:
                            for d3 in d2['children']:
                                if d3['label'] != "":
                                    case_names_three.append(d3['label'])
                                    if 'children' in d3:
                                        for d4 in d3['children']:
                                            if d4['label'] != "":
                                                case_names_four.append(d4['label'])
                                                if 'children' in d4:
                                                    for d5 in d4['children']:
                                                        if d5['label'] != "":
                                                            case_names_five.append(d5['label'])

    return case_names_one, case_names_two, case_names_three, case_names_four, case_names_five


def read_data(path, save_path, case_name):
    document = docx.Document(path)
    # 提取指定章节
    chapter_title = case_name
    case_name = re.sub(r"[\/\\\:\*\?\"\<\>\|]", "", case_name)
    output_filename = os.path.join(save_path, case_name + ".docx")
    extract_chapter_to_new_docx(document, chapter_title, output_filename)


if __name__ == '__main__':
    input_path = 'D:\\PyProject\\pythonProject\\功能规格合并版.docx'
    # document = docx.Document(input_path)
    # json_structure = convert_to_json_structure(document)
    # print(json_structure)
    one, two, three, four, five = read_title(input_path)
    read_data(input_path, "D:\\PyProject\\pythonProject", '工艺设计详细方案')
